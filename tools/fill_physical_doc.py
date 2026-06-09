from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"d:\database")
OUT = ROOT / "docs" / "3 物理结构设计&系统实现.docx"
BACKUP = ROOT / "docs" / "3 物理结构设计&系统实现.backup.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_code(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11, "1F4E79"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)


def para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def build():
    if OUT.exists() and not BACKUP.exists():
        BACKUP.write_bytes(OUT.read_bytes())

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    r = title.add_run("物理结构设计&系统实现")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)

    for line in [
        "课题名称：校园活动管理系统设计",
        "课题组员：王熙泽、蓝莹、唐瑀含、何佳倩",
        "完成时间：2026年6月7日",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)
    doc.add_page_break()

    doc.add_heading("6 物理设计", level=1)
    doc.add_heading("6.1 物理设计目标", level=2)
    para(doc, "校园活动管理系统采用 B/S 架构，数据库端使用 SQL Server Express，后端使用 Spring Boot 3 与 MyBatis-Plus 访问数据库，前端使用 Vue 3、Vite、Element Plus 实现三类角色工作台。物理设计的目标是在保证数据完整性和演示可运行性的基础上，提高活动查询、报名名单查询、签到统计和评分统计的访问效率。")
    add_bullets(doc, [
        "保证学生、组织者、管理员、活动、报名、签到、评分七类核心数据能够稳定存储。",
        "通过主键、外键、唯一约束和 CHECK 约束保证基础完整性。",
        "通过组合索引支持活动列表筛选、组织者活动管理、学生报名历史和统计查询。",
        "通过视图、存储过程和触发器补充课程设计要求中的 SQL 实现内容。",
    ])

    doc.add_heading("6.2 数据库运行环境", level=2)
    add_table(doc, ["项目", "当前实现"], [
        ["数据库系统", "Microsoft SQL Server Express，本机实例 .\\SQLEXPRESS"],
        ["数据库名称", "campus_activity"],
        ["连接方式", "Windows 集成认证；后端 start.ps1 自动读取 SQLEXPRESS 动态 TCP 端口"],
        ["后端数据访问", "mssql-jdbc + MyBatis-Plus BaseMapper"],
        ["初始化脚本", "sql/init.sql，包含建库、建表、约束、索引、测试数据、视图、存储过程、触发器"],
    ])

    doc.add_heading("6.3 数据表物理结构", level=2)
    add_table(doc, ["表名", "主要字段", "物理设计说明"], [
        ["student", "student_id, student_no, password, student_name, college, major, grade, phone, email, interest_tags, account_status", "student_id 使用 BIGINT IDENTITY 主键；student_no 唯一；姓名、学院等使用 NVARCHAR 保存文本。"],
        ["organizer", "organizer_id, organizer_no, password, organizer_name, organizer_type, contact_name, contact_phone, account_status", "organizer_id 使用自增主键；organizer_no 唯一；用于区分社团、学院部门、校级组织。"],
        ["admin_user", "admin_id, username, password, admin_name, account_status", "用于基础管理员登录与账号状态管理。"],
        ["activity", "activity_id, activity_title, category, organizer_id, location, start_time, end_time, register_start_time, register_end_time, capacity, status", "保存活动发布信息；时间字段使用 DATETIME2；organizer_id 外键引用 organizer。"],
        ["registration", "registration_id, activity_id, student_id, registration_time, registration_status, cancel_time", "保存报名流水；student_id + activity_id 唯一，防止重复报名。"],
        ["checkin", "checkin_id, registration_id, activity_id, student_id, checkin_time, checkin_method, checkin_status, remark", "保存签到或补签记录；registration_id 唯一，防止重复签到。"],
        ["rating", "rating_id, activity_id, student_id, checkin_id, score, comment, rating_time", "保存评分反馈；score 使用 TINYINT 并限制 1 到 5 分。"],
    ])

    doc.add_heading("6.4 索引设计", level=2)
    add_table(doc, ["索引名称", "所在表", "字段", "设计目的"], [
        ["ix_activity_status_time", "activity", "audit_status, activity_status, register_end_time, start_time", "支持学生端活动列表、推荐活动过滤和按时间排序。"],
        ["ix_activity_organizer_time", "activity", "organizer_id, start_time", "支持组织者查看自己发布的活动。"],
        ["ix_registration_activity_status", "registration", "activity_id, registration_status", "支持组织者查看报名名单和统计报名人数。"],
        ["ix_registration_student_status", "registration", "student_id, registration_status", "支持学生查看我的报名、推荐过滤已报名活动。"],
        ["ix_checkin_activity_student", "checkin", "activity_id, student_id", "支持签到校验和签到人数统计。"],
        ["ix_rating_activity_score", "rating", "activity_id, score", "支持平均评分、评分人数统计。"],
    ])
    add_code(doc, """
CREATE INDEX ix_activity_status_time
ON activity (audit_status, activity_status, register_end_time, start_time);
CREATE INDEX ix_activity_organizer_time ON activity (organizer_id, start_time);
CREATE INDEX ix_registration_activity_status ON registration (activity_id, registration_status);
CREATE INDEX ix_registration_student_status ON registration (student_id, registration_status);
CREATE INDEX ix_checkin_activity_student ON checkin (activity_id, student_id);
CREATE INDEX ix_rating_activity_score ON rating (activity_id, score);
""")

    doc.add_heading("7 系统SQL实现", level=1)
    doc.add_heading("7.1 完整性设计", level=2)
    para(doc, "系统完整性设计包括实体完整性、参照完整性、用户定义完整性和业务完整性四类。实体完整性通过各表自增主键保证；参照完整性通过外键保证；用户定义完整性通过唯一约束、默认值和 CHECK 约束保证；业务完整性通过后端 Service 校验与数据库触发器共同保证。")
    add_table(doc, ["完整性类型", "实现方式", "说明"], [
        ["实体完整性", "PRIMARY KEY + BIGINT IDENTITY", "七张核心表均使用自增主键。"],
        ["参照完整性", "FOREIGN KEY", "活动必须属于组织者；报名必须关联学生和活动；签到必须关联报名；评分必须关联签到。"],
        ["唯一性", "UNIQUE", "student_no、organizer_no、admin username 唯一；同一学生同一活动只能报名、评分一次。"],
        ["范围约束", "CHECK", "活动结束时间晚于开始时间；报名截止不晚于活动开始；capacity > 0；score BETWEEN 1 AND 5。"],
        ["业务约束", "触发器 + 后端校验", "报名人数不能超过活动容量，后端也会校验时间、名额、重复报名、重复签到和重复评分。"],
    ])

    doc.add_heading("7.2 表的创建", level=2)
    para(doc, "完整建表脚本位于 sql/init.sql。核心建表语句节选如下，覆盖优化后的七个关系模式。")
    add_code(doc, """
CREATE TABLE student (
    student_id BIGINT IDENTITY(1,1) PRIMARY KEY,
    student_no VARCHAR(30) NOT NULL UNIQUE,
    password VARCHAR(100) NOT NULL,
    student_name NVARCHAR(50) NOT NULL,
    account_status VARCHAR(20) NOT NULL DEFAULT 'enabled',
    created_time DATETIME2 NOT NULL DEFAULT SYSDATETIME()
);

CREATE TABLE activity (
    activity_id BIGINT IDENTITY(1,1) PRIMARY KEY,
    activity_title NVARCHAR(100) NOT NULL,
    category NVARCHAR(50) NOT NULL,
    organizer_id BIGINT NOT NULL,
    start_time DATETIME2 NOT NULL,
    end_time DATETIME2 NOT NULL,
    register_start_time DATETIME2 NOT NULL,
    register_end_time DATETIME2 NOT NULL,
    capacity INT NOT NULL,
    activity_status VARCHAR(20) NOT NULL DEFAULT 'draft',
    audit_status VARCHAR(20) NOT NULL DEFAULT 'pending',
    CONSTRAINT fk_activity_organizer FOREIGN KEY (organizer_id) REFERENCES organizer(organizer_id),
    CONSTRAINT ck_activity_time CHECK (end_time > start_time),
    CONSTRAINT ck_activity_register_time CHECK (register_end_time <= start_time),
    CONSTRAINT ck_activity_capacity CHECK (capacity > 0)
);
""")

    doc.add_heading("7.3 安全性设计", level=2)
    para(doc, "数据库层当前采用 SQL Server Windows 集成认证，避免在配置文件中固定数据库账号密码。应用层使用简单 Token 进行会话识别，登录后前端将 token 保存到 localStorage，并在请求头 Authorization 中传给后端。后端依据 AuthUser.role 区分 student、organizer、admin 三类角色。")
    add_table(doc, ["角色", "可访问功能", "后端控制方式"], [
        ["学生 student", "活动浏览、报名、取消报名、签到、评分、推荐活动", "RegistrationService、CheckinService、RatingService 中调用 requireRole。"],
        ["组织者 organizer", "活动创建/发布/取消、报名名单、签到码、人工补签、统计", "ActivityService.assertOrganizerOwner 限制只能管理自己的活动。"],
        ["管理员 admin", "用户状态管理、活动审核", "AdminService 与 ActivityService.audit 中限制 admin 角色。"],
    ])
    add_code(doc, """
-- 数据库角色设计示例，可在正式部署时启用
CREATE ROLE campus_reader;
CREATE ROLE campus_operator;
CREATE ROLE campus_admin;
GRANT SELECT ON vw_activity_summary TO campus_reader;
GRANT SELECT, INSERT, UPDATE ON registration TO campus_operator;
GRANT SELECT, INSERT, UPDATE ON activity TO campus_operator;
GRANT CONTROL ON DATABASE::campus_activity TO campus_admin;
""")

    doc.add_heading("7.4 视图的创建", level=2)
    para(doc, "系统创建两个视图：vw_activity_summary 用于活动统计汇总，vw_student_participation 用于查询学生参与历史。它们可以减少前端和后端联表统计的复杂度，也便于演示数据库查询能力。")
    add_code(doc, """
CREATE VIEW vw_activity_summary AS
SELECT a.activity_id, a.activity_title, a.category, a.location,
       a.start_time, a.end_time, a.capacity, a.activity_status, a.audit_status,
       o.organizer_name,
       COUNT(DISTINCT CASE WHEN r.registration_status = 'registered' THEN r.registration_id END) AS registration_count,
       COUNT(DISTINCT c.checkin_id) AS checkin_count,
       CAST(ISNULL(AVG(CAST(rt.score AS DECIMAL(6,2))), 0) AS DECIMAL(6,2)) AS average_score
FROM activity a
JOIN organizer o ON a.organizer_id = o.organizer_id
LEFT JOIN registration r ON a.activity_id = r.activity_id
LEFT JOIN checkin c ON a.activity_id = c.activity_id
LEFT JOIN rating rt ON a.activity_id = rt.activity_id
GROUP BY a.activity_id, a.activity_title, a.category, a.location,
         a.start_time, a.end_time, a.capacity, a.activity_status,
         a.audit_status, o.organizer_name;

-- 演示
SELECT * FROM vw_activity_summary WHERE activity_id = 5;
""")

    doc.add_heading("7.5 存储过程的创建", level=2)
    para(doc, "系统创建 usp_activity_statistics 与 usp_register_activity 两个存储过程。前者用于按活动编号输出报名人数、签到人数、签到率和平均评分；后者演示数据库层报名校验流程。")
    add_code(doc, """
CREATE PROCEDURE usp_activity_statistics
    @activity_id BIGINT
AS
BEGIN
    SET NOCOUNT ON;
    SELECT activity_id, activity_title, registration_count,
           checkin_count, checkin_rate, average_score
    FROM vw_activity_summary
    WHERE activity_id = @activity_id;
END;

-- 演示
EXEC usp_activity_statistics 5;
""")
    add_code(doc, """
CREATE PROCEDURE usp_register_activity
    @student_id BIGINT,
    @activity_id BIGINT
AS
BEGIN
    SET NOCOUNT ON;
    -- 校验学生状态、活动报名时间、名额和重复报名后插入 registration。
    -- 完整脚本见 sql/init.sql。
END;
""")

    doc.add_heading("7.6 触发器的创建", level=2)
    para(doc, "为了在数据库层防止活动报名人数超过 capacity，系统创建 tr_registration_capacity_guard 触发器。该触发器在 registration 表 INSERT 或 UPDATE 后检查每个活动的有效报名数，如果超过 capacity 则回滚事务。")
    add_code(doc, """
CREATE TRIGGER tr_registration_capacity_guard
ON registration
AFTER INSERT, UPDATE
AS
BEGIN
    SET NOCOUNT ON;
    IF EXISTS (
        SELECT 1
        FROM activity a
        JOIN (
            SELECT activity_id, COUNT(*) AS registered_count
            FROM registration
            WHERE registration_status = 'registered'
            GROUP BY activity_id
        ) x ON a.activity_id = x.activity_id
        WHERE x.registered_count > a.capacity
    )
    BEGIN
        ROLLBACK TRANSACTION;
        THROW 50005, 'Registration count cannot exceed activity capacity.', 1;
    END;
END;
""")

    doc.add_heading("8 系统实现", level=1)
    doc.add_heading("8.1 后端实现", level=2)
    para(doc, "后端位于 backend 目录，采用 Spring Boot 3、Java 21、MyBatis-Plus 和 SQL Server JDBC Driver。代码按 controller、service、mapper、entity、dto、vo、common 分层组织。统一返回格式为 { code, message, data }。")
    add_table(doc, ["层次", "主要文件", "职责"], [
        ["Controller", "AuthController、ActivityController、RegistrationController、CheckinController、RatingController、StatisticsController、RecommendationController、AdminController", "接收 HTTP 请求并返回统一 ApiResponse。"],
        ["Service", "AuthService、ActivityService、RegistrationService、CheckinService、RatingService、StatisticsService、RecommendationService、AdminService", "实现登录、权限判断和核心业务规则。"],
        ["Mapper", "StudentMapper、ActivityMapper 等", "继承 MyBatis-Plus BaseMapper 完成数据库访问。"],
        ["Entity/DTO/VO", "Student、Activity、LoginRequest、StatisticsVO 等", "对应数据库表、请求参数和接口返回对象。"],
    ])

    doc.add_heading("8.2 前端实现", level=2)
    para(doc, "前端位于 frontend 目录，采用 Vue 3、Vite、Element Plus、Axios、Vue Router 和 Pinia。界面已套用简约后台模板风格，包含登录页、深色侧边栏、顶部用户栏和三类角色工作台。")
    add_table(doc, ["角色", "页面", "功能"], [
        ["学生", "ActivityListView、ActivityDetailView、MyRegistrationsView、StudentCheckinView、StudentRatingView、RecommendationsView", "活动浏览、报名、取消报名、签到、评分、推荐活动。"],
        ["组织者", "OrganizerActivitiesView、ActivityFormView、RegistrationListView、OrganizerCheckinView、StatisticsView", "活动发布与管理、报名名单、签到码、人工补签、活动统计。"],
        ["管理员", "AdminUsersView、AdminActivitiesView", "用户状态管理、活动审核。"],
    ])

    doc.add_heading("8.3 接口实现", level=2)
    para(doc, "后端接口以 REST 风格提供，接口文档位于 docs/api.md。核心接口包括登录、活动、报名、签到、评分、统计和推荐。")
    add_table(doc, ["模块", "接口示例", "说明"], [
        ["认证", "POST /api/login；GET /api/users/me", "三类角色登录和当前用户查询。"],
        ["活动", "GET /api/activities；POST /api/activities；PUT /api/activities/{id}/publish", "活动列表、详情、创建、修改、发布、取消。"],
        ["报名", "POST /api/registrations；GET /api/registrations/mine", "学生报名、取消报名、我的报名、报名名单。"],
        ["签到评分", "POST /api/checkins/code；POST /api/checkins；POST /api/ratings", "生成签到码、学生签到、人工补签、活动评分。"],
        ["统计推荐", "GET /api/statistics/activity/{id}；GET /api/recommendations", "活动统计和规则推荐。"],
        ["管理员", "GET /api/admin/users；PUT /api/admin/activities/{id}/audit", "用户状态和活动审核。"],
    ])

    doc.add_heading("8.4 运行与验证", level=2)
    para(doc, "当前系统已经完成数据库初始化、后端编译、前端构建和接口冒烟测试。前端运行地址为 http://localhost:5173/，后端运行地址为 http://localhost:8080/。")
    add_table(doc, ["验证项", "结果"], [
        ["数据库初始化", "sql/init.sql 可执行，创建 7 张表、6 个索引、2 个视图、2 个存储过程、1 个触发器，并插入测试数据。"],
        ["后端构建", "mvn -q -DskipTests package 通过。"],
        ["前端构建", "npm run build 通过。"],
        ["接口测试", "学生登录、管理员登录、活动列表、我的报名、推荐活动、活动统计均返回成功。"],
        ["演示账号", "学生 2026001/123456；组织者 org001/123456；管理员 admin/123456。"],
    ])

    doc.save(OUT)


if __name__ == "__main__":
    build()
